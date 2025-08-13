# newsletter_v3.py
import streamlit as st
from io import BytesIO
from datetime import datetime
import re

# ---- ReportLab (PDF) ----
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image, ListFlowable, ListItem, KeepTogether, Table, TableStyle
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib.colors import lightgrey, Color, black, HexColor
from reportlab.lib.utils import ImageReader

# ---- python-docx (DOCX) ----
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---- PIL for image overlay (month/year on top image) ----
from PIL import Image as PILImage, ImageDraw, ImageFont

# =========================================================
# Helpers
# =========================================================
def process_content_pdf(content: str) -> str:
    """
    Convert [text](url) to <a> for PDF; also auto-link urls/emails.
    """
    if not content:
        return ""
    # Convert markdown-style links to <a>
    processed = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2" color="blue">\1</a>', content)

    # Auto-link plain URLs and emails (avoid double-wrapping by excluding content inside <a> tags)
    url_pattern = re.compile(
        r'(?<!href=")(?!.*</a>)(https?://\S+|mailto:\S+|\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})(?!")(?![^<]*</a>)'
    )

    def repl(m):
        link = m.group(0)
        if '@' in link and not link.startswith('mailto:'):
            link = f'mailto:{link}'
        return f'<a href="{link}" color="blue">{m.group(0)}</a>'

    return re.sub(url_pattern, repl, processed)


def add_hyperlink(paragraph, url, text=None):
    """
    Insert a clickable hyperlink into a python-docx paragraph.
    """
    if not text:
        text = url
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.set(qn('w:history'), '1')

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Color blue
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run_text = OxmlElement('w:t')
    new_run_text.text = text
    new_run.append(new_run_text)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph


def add_text_with_links(paragraph, text: str):
    """
    python-docx helper to render [text](url) and raw links/emails as clickable.
    """
    if not text:
        return
    md = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
    last_end = 0
    for m in md.finditer(text):
        paragraph.add_run(text[last_end:m.start()])
        add_hyperlink(paragraph, m.group(2), m.group(1))
        last_end = m.end()
    tail = text[last_end:]

    url_pattern = re.compile(r'(https?://\S+|mailto:\S+|\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})')
    last_end = 0
    for m in url_pattern.finditer(tail):
        paragraph.add_run(tail[last_end:m.start()])
        link = m.group(0)
        if '@' in link and not link.startswith('mailto:'):
            link = f'mailto:{link}'
        add_hyperlink(paragraph, link, m.group(0))
        last_end = m.end()
    paragraph.add_run(tail[last_end:])


def text_to_bullets(text: str):
    if not text:
        return []
    return [line.strip() for line in text.splitlines() if line.strip()]


COLOR_OPTIONS = {
    "Cream White": Color(0.98, 0.96, 0.92),
    "Off White": Color(0.96, 0.94, 0.90),
    "Light Blue": Color(0.90, 0.95, 1.0),
    "Light Green": Color(0.90, 0.98, 0.90),
    "Light Pink": Color(0.98, 0.90, 0.95),
    "Light Yellow": Color(0.98, 0.98, 0.85)
}

# =========================================================
# Page decorations for PDF
# =========================================================
def draw_page_frame(canv, doc, headers, footers, draw_confidential=False, confidentiality_line="", draw_watermark=False):
    """
    Draws horizontal lines (top/bottom), header/footer (left/right),
    optional confidentiality (first page only), optional diagonal watermark.
    """
    width, height = A4

    # Horizontal lines
    canv.saveState()
    canv.setStrokeColor(black)
    canv.setLineWidth(0.5)
    canv.line(30, height - 40, width - 30, height - 40)  # top line
    canv.line(30, 40, width - 30, 40)                    # bottom line

    # Header/Footer text
    canv.setFont("Helvetica", 9)
    # Header
    if headers.get('top_left'):
        canv.drawString(40, height - 30, headers['top_left'])
    if headers.get('top_right'):
        canv.drawRightString(width - 40, height - 30, headers['top_right'])
    # Footer
    if footers.get('bottom_left'):
        canv.drawString(40, 30, footers['bottom_left'])
    if footers.get('bottom_right'):
        canv.drawRightString(width - 40, 30, footers['bottom_right'])

    # Confidentiality (first page only, centered near bottom)
    if draw_confidential and confidentiality_line:
        canv.setFont("Helvetica-Oblique", 8)
        canv.drawCentredString(width / 2.0, 50, confidentiality_line)

    # Optional watermark across page
    if draw_watermark:
        canv.setFont("Helvetica", 48)
        canv.setFillGray(0.9, 0.3)  # light watermark
        canv.saveState()
        canv.translate(width / 2.0, height / 2.0)
        canv.rotate(45)
        canv.drawCentredString(0, 0, "NEWSLETTER")
        canv.restoreState()

    canv.restoreState()


def build_pdf(story, headers, footers, confidentiality_line, add_watermark):
    """
    Build the PDF with consistent frames on all pages and confidentiality line on first page only.
    """
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=50, rightMargin=50, topMargin=70, bottomMargin=70)

    def _first(canv, doc_):
        draw_page_frame(
            canv, doc_, headers, footers,
            draw_confidential=True,
            confidentiality_line=confidentiality_line,
            draw_watermark=add_watermark
        )

    def _later(canv, doc_):
        draw_page_frame(
            canv, doc_, headers, footers,
            draw_confidential=False,
            confidentiality_line="",
            draw_watermark=add_watermark
        )

    doc.build(story, onFirstPage=_first, onLaterPages=_later)
    buf.seek(0)
    return buf

# =========================================================
# Image utilities
# =========================================================
def overlay_month_year_on_image(file, month: str, year: str):
    """
    Overlays 'Month Year' text on the top-right of the uploaded image.
    Returns a BytesIO object for ReportLab Image.
    """
    try:
        img = PILImage.open(file).convert("RGBA")
        draw = ImageDraw.Draw(img)

        text = f"{month} {year}".strip()
        if not text:
            bio = BytesIO()
            img.save(bio, format="PNG")
            bio.seek(0)
            return bio

        # Choose a font (default if no TTF available)
        try:
            font = ImageFont.truetype("arial.ttf", 28)
        except:
            font = ImageFont.load_default()

        text_w, text_h = draw.textbbox((0, 0), text, font=font)[2:]
        pad = 10
        x = img.width - text_w - 2 * pad - 10
        y = 10
        # semi-opaque white box behind text for readability
        box = (x, y, x + text_w + 2 * pad, y + text_h + 2 * pad)
        draw.rectangle(box, fill=(255, 255, 255, 180))
        draw.text((x + pad, y + pad), text, fill=(0, 0, 0, 255), font=font)

        out = BytesIO()
        img.save(out, format="PNG")
        out.seek(0)
        return out
    except Exception:
        # If anything goes wrong, just return original file
        if hasattr(file, "seek"):
            file.seek(0)
        raw = BytesIO(file.read() if hasattr(file, "read") else file)
        raw.seek(0)
        return raw

# =========================================================
# PDF Generation
# =========================================================
def create_pdf(top_image_file, intro_text, intro_color, sections, contact_info, bottom_image_file,
               headers, footers, confidentiality_line, add_watermark):
    styles = getSampleStyleSheet()
    story = []
    
    # Calculate content width (A4 width - left margin - right margin)
    content_width = A4[0] - 100  # 50 left + 50 right margin = 100

    # Top image
    if top_image_file is not None:
        top_image_file.seek(0)
        img_io = BytesIO(top_image_file.read())
        img_io.seek(0)
        story.append(Image(img_io, width=content_width, height=2 * inch))
        story.append(Spacer(1, 0.25 * inch))

    # Intro block with colored background
    if intro_text:
        intro_style = ParagraphStyle(
            "Intro",
            parent=styles["Normal"],
            backColor=intro_color,
            leading=14,
            spaceAfter=12,
            borderPadding=10
        )
        story.append(Paragraph(process_content_pdf(intro_text), intro_style))
        story.append(Spacer(1, 0.2 * inch))

    # Sections
    for i, sec in enumerate(sections, start=1):
        title = (sec.get("title") or "").strip()
        content = (sec.get("content") or "").strip()
        section_color = sec.get("color", Color(1, 1, 1))

        bullets = text_to_bullets(content)
        if title or bullets:
            section_content = []
            
            if title:
                section_content.append(Paragraph(f"<b>{title}</b>", styles["Heading2"]))
                section_content.append(Spacer(1, 0.05 * inch))
            
            if bullets:
                bullet_style = ParagraphStyle(f"BulletStyle{i}", parent=styles["Normal"], leading=14, spaceBefore=0, spaceAfter=28, bulletFontSize=8)
                items = [ListItem(Paragraph(process_content_pdf(b), bullet_style), leftIndent=12) for b in bullets]
                section_content.append(ListFlowable(items, bulletType='bullet', start='circle', leftIndent=18, bulletFontSize=8))
            
            # Create single paragraph with section background
            if title and bullets:
                combined_text = f"<b>{title}</b><br/><br/>" + "<br/><br/>".join([f"• {process_content_pdf(b)}" for b in bullets])
            elif title:
                combined_text = f"<b>{title}</b>"
            else:
                combined_text = "<br/><br/>".join([f"• {process_content_pdf(b)}" for b in bullets])
            
            section_style = ParagraphStyle(f"Section{i}", parent=styles["Normal"], leading=14, spaceAfter=12, backColor=section_color, borderPadding=10)
            story.append(Paragraph(combined_text, section_style))
            # Use 2-line spacing for last 4 sections
            spacing = 0.3 * inch if i > len(sections) - 4 else 0.15 * inch
            story.append(Spacer(1, spacing))

    # Contact info
    if contact_info:
        story.append(Paragraph("<b>Contact Information</b>", styles["Heading2"]))
        story.append(Spacer(1, 0.05 * inch))
        contact_style = ParagraphStyle("ContactStyle", parent=styles["Normal"], leading=14, spaceAfter=12)
        story.append(Paragraph(process_content_pdf(contact_info), contact_style))

    # Bottom image
    if bottom_image_file is not None:
        bottom_image_file.seek(0)
        bio = BytesIO(bottom_image_file.read())
        bio.seek(0)
        story.append(Spacer(1, 0.3 * inch))
        story.append(Image(bio, width=content_width, height=1 * inch))

    return build_pdf(
        story=story,
        headers=headers,
        footers=footers,
        confidentiality_line=confidentiality_line,
        add_watermark=add_watermark
    )

# =========================================================
# DOCX Generation
# =========================================================
def create_docx(top_image_file, intro_text, sections, contact_info, bottom_image_file,
                headers, footers):
    doc = Document()
    section = doc.sections[0]
    
    # Calculate content width (A4 width - margins)
    content_width_inches = (A4[0] - 100) / 72  # Convert points to inches

    # Header (2 cells: left/right)
    header = section.header
    h_para = header.add_paragraph()
    h_para.add_run(headers.get('top_left', '') or '').bold = True
    h_para.add_run('\t' + (headers.get('top_right', '') or '')).bold = True

    # Footer (2 cells: left/right)
    footer = section.footer
    f_para = footer.add_paragraph()
    f_para.add_run(footers.get('bottom_left', '') or '')
    f_para.add_run('\t' + (footers.get('bottom_right', '') or ''))

    # Top image
    if top_image_file is not None:
        p = doc.add_paragraph()
        run = p.add_run()
        top_image_file.seek(0)
        run.add_picture(top_image_file, width=Inches(content_width_inches), height=Inches(2))

    # Intro text
    if intro_text:
        p = doc.add_paragraph()
        add_text_with_links(p, intro_text)

    # Sections
    for sec in sections:
        title = (sec.get("title") or "").strip()
        content = (sec.get("content") or "").strip()
        if title:
            doc.add_heading(title, level=2)
        for bullet in text_to_bullets(content):
            p = doc.add_paragraph(style='List Bullet')
            add_text_with_links(p, bullet)

    # Contact info
    if contact_info:
        doc.add_heading("Contact Information", level=2)
        p = doc.add_paragraph()
        add_text_with_links(p, contact_info)

    # Bottom image
    if bottom_image_file is not None:
        p = doc.add_paragraph()
        run = p.add_run()
        bottom_image_file.seek(0)
        run.add_picture(bottom_image_file, width=Inches(content_width_inches), height=Inches(1))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# =========================================================
# Streamlit UI
# =========================================================
st.set_page_config(page_title="Newsletter Generator", page_icon="📬", layout="centered")

st.title("📬 Newsletter Generator")
st.caption("Multi-section newsletter with images, links, headers/footers, and a first-page confidentiality note.")

with st.expander("Images"):
    top_image = st.file_uploader("Top image (optional)", type=["png", "jpg", "jpeg", "webp"])
    bottom_image = st.file_uploader("Bottom image (optional)", type=["png", "jpg", "jpeg", "webp"])



st.subheader("Intro")
intro_text = st.text_area("Intro text (links supported: [label](https://example.com))", height=120)
intro_color_name = st.selectbox("Intro background color", list(COLOR_OPTIONS.keys()), index=0)
intro_color = COLOR_OPTIONS[intro_color_name]

st.subheader("Sections")
num_sections = st.number_input("How many sections?", min_value=1, max_value=10, value=3, step=1)
sections = []
for i in range(int(num_sections)):
    st.markdown(f"**Section {i+1}**")
    title = st.text_input(f"Title {i+1}", key=f"title_{i}")
    content = st.text_area(
        f"Bullets {i+1} (one per line; links supported)", key=f"content_{i}", height=120
    )
    color_name = st.selectbox(f"Section {i+1} background color", list(COLOR_OPTIONS.keys()), index=0, key=f"color_{i}")
    sections.append({"title": title, "content": content, "color": COLOR_OPTIONS[color_name]})

st.subheader("Contact")
contact_info = st.text_area("Contact info (links supported)", height=100)

st.subheader("Headers & Footers (all pages)")
col1, col2 = st.columns(2)
with col1:
    header_top_left = st.text_input("Header — Top Left", placeholder="Company Name")
    footer_bottom_left = st.text_input("Footer — Bottom Left", placeholder="Address or tagline")
with col2:
    header_top_right = st.text_input("Header — Top Right", placeholder="www.example.com")
    footer_bottom_right = st.text_input("Footer — Bottom Right", placeholder="Page/Legal/etc.")

st.subheader("Confidentiality (first PDF page only)")
confidentiality_line = st.text_input("Confidentiality line (centered at bottom of 1st page)")

add_watermark = st.checkbox("Add diagonal 'NEWSLETTER' watermark", value=False)

st.markdown("---")
if st.button("Generate PDF and DOCX"):
    headers = {"top_left": header_top_left, "top_right": header_top_right}
    footers = {"bottom_left": footer_bottom_left, "bottom_right": footer_bottom_right}

    # Create PDF + DOCX
    pdf_buffer = create_pdf(
        top_image, intro_text, intro_color, sections, contact_info, bottom_image,
        headers, footers, confidentiality_line, add_watermark
    )
    docx_buffer = create_docx(
        top_image, intro_text, sections, contact_info, bottom_image,
        headers, footers
    )

    st.success("✅ Newsletter generated successfully!")
    dl1, dl2 = st.columns(2)
    with dl1:
        st.download_button(
            "📄 Download PDF",
            data=pdf_buffer,
            file_name="newsletter.pdf",
            mime="application/pdf",
            use_container_width=True
        )
    with dl2:
        st.download_button(
            "📝 Download DOCX",
            data=docx_buffer,
            file_name="newsletter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
